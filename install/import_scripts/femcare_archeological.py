from __future__ import annotations

import re
import shutil
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Sequence

import pandas as pd
from docx import Document

from openatlas import app
from openatlas.api.import_scripts.util import get_exact_match
from openatlas.models.entity import Entity, insert

FILE_PATH = Path('files/femcare')
SKELETON_PATH = FILE_PATH / 'Elisabethinen_SE_Protokolle'
SCHNITTE_PATH = FILE_PATH / "Schnitte"
SKELETON_IMAGE_PATH = FILE_PATH / 'skelett_mannchen'
UPLOAD = FILE_PATH / 'uploads'
FUNDFOTOS = FILE_PATH / 'Fundfotos'
FUNDFOTOS_KATALOG = FILE_PATH / 'Katalog_fundfotos.docx'
FUNDFOTOS_MEDAILLIEN = FILE_PATH / 'Fundfotos_Medaillien'

# pylint: skip-file

DEBUG_MSG = defaultdict(list)

FILE_INFO = {
    'creator': 'LH Novetus',
    'license_holder': 'LH Novetus',
    'public': True}


@dataclass
class Individual:
    se_id: int
    description: list[str]
    probe_type: str
    # find_type: str
    position: str
    orientation: str
    preservation: str
    dislocation: str
    age: str
    extraction: str


@dataclass
class ParsedStratigraphicUnit:
    id_: str
    se_id: int
    type: str
    name: str
    layer: str
    feature: str
    individual_id: Optional[int] = None


@dataclass
class Feature:
    id_: str
    obj_id: int
    name: str
    type: str
    se: list[str]
    description: str | None
    cut: str


@dataclass
class Find:
    id_: str
    f_id: int
    name: str
    material: str
    designation: str
    description: str
    dating: str
    stratigraphic_unit: str
    feature_id: str
    openatlas_class: str


@dataclass
class StratigraphicUnit:
    id_: str
    se_id: int
    name: str
    layer: str
    feature: str
    type: str
    description: Optional[list[str]] = None
    individual_id: Optional[int] = None
    probe_type: Optional[str] = None
    find_type: Optional[str] = None
    position: Optional[str] = None
    orientation: Optional[str] = None
    preservation: Optional[str] = None
    dislocation: Optional[str] = None
    age: Optional[str] = None
    extraction: Optional[str] = None


@dataclass
class FundEntity:
    id: str
    type: str = ''
    date: str = ''
    location: str = ''
    description: str = ''
    dimensions: str = ''
    identifications: str = ''

    _current_index: int = field(default=0, init=False, repr=False)

    def add_entry(self, text: str) -> None:
        field_names = [
            'type',
            'date',
            'location',
            'description',
            'dimensions',
            'identifications']
        if self._current_index < len(field_names):
            target_field = field_names[self._current_index]
            setattr(self, target_field, text)
            self._current_index += 1


@dataclass
class KatalogEntity:
    id: str
    type: str = 'Medaille'
    start_date: str = ''
    end_date: str = ''
    location: str = ''
    description: str = ''
    weight: str = ''
    diameter: str = ''
    coin: str = ''
    length: str = ''
    height: str = ''
    identifications: str = ''
    fndnr: str = ''
    se: str = ''
    image_id: str = ''


def build_se_ind_map(root: Path) -> dict[str, Path]:
    out: dict[str, Path] = {}
    for path in root.rglob("*"):
        if not path.is_dir():
            continue
        name = path.name
        if name.startswith("SE_") and "_Ind_" in name:
            out[name] = path
    return out


def parse_features() -> list[Feature]:
    df = pd.read_csv(FILE_PATH / 'features.csv', delimiter=',')
    features_ = []
    current_cut = ''
    for _, row in df.iterrows():
        if "Schnitt" in row.iloc[0]:
            current_cut = row.iloc[0]
            continue
        se_raw = row.iloc[2]
        se_list = se_raw.split(", ") if pd.notna(se_raw) else []
        entry_obj = Feature(
            id_=f"feature_{int(row.iloc[0])}".strip(),
            obj_id=int(row.iloc[0]),
            name=f'Objekt {str(int(row.iloc[0])).strip()}',
            type=row.iloc[1].strip(),
            se=se_list,
            description=row.iloc[3] if pd.notna(row.iloc[3]) else None,
            cut=current_cut.strip())
        features_.append(entry_obj)
    return features_


def parse_stratigraphic_units() -> list[ParsedStratigraphicUnit]:
    df = pd.read_csv(FILE_PATH / 'se.csv', delimiter=',')
    se = []
    for _, row in df.iterrows():
        if pd.isna(row.iloc[2]) or row.iloc[2] == '':
            continue
        if pd.isna(row.iloc[4]) or row.iloc[4] == 0:
            DEBUG_MSG['no_feature_available'].append(int(row.iloc[0]))
            continue
        entry_obj = ParsedStratigraphicUnit(
            id_=f"stratigraphic_{int(row.iloc[0])}".strip(),
            se_id=int(row.iloc[0]),
            name=f'SE {str(int(row.iloc[0])).strip()}',
            # convert to int cause of the leading zeros
            type=row.iloc[1],
            layer=row.iloc[2].strip(),
            individual_id=int(row.iloc[3]) if pd.notna(row.iloc[3]) else None,
            feature=f"feature_{int(row.iloc[4])}".strip()
            if pd.notna(row.iloc[4]) else '')
        se.append(entry_obj)
    return se


def parse_finds() -> list[Find]:
    df = pd.read_csv(FILE_PATH / 'finds.csv', delimiter=',')
    finds_ = []
    for _, row in df.iterrows():

        stratigraphic_unit = ''
        if pd.notna(row.iloc[1]):
            if row.iloc[1] == '-':
                stratigraphic_unit = ''
            elif '/' in row.iloc[1]:
                DEBUG_MSG['multiple_SE_in_finds'].append(row.iloc[0])
            elif '?' in row.iloc[1]:
                DEBUG_MSG['undecided_SE_in_finds'].append(row.iloc[0])
            else:
                stratigraphic_unit = (f"stratigraphic_"
                                      f"{int(row.iloc[1])}").strip()
        feature_ = ''
        if pd.notna(row.iloc[4]):
            if row.iloc[4] == '-':
                feature_ = ''
            else:
                feature_ = f"feature_{int(row.iloc[4])}".strip()

        entry_obj = Find(
            id_=f"find_{int(row.iloc[0])}".strip(),
            f_id=int(row.iloc[0]),
            stratigraphic_unit=stratigraphic_unit,
            feature_id=feature_,
            name=f"FNR {int(row.iloc[0])}",
            material=row.iloc[6] if pd.notna(row.iloc[6]) else '',
            designation=row.iloc[7] if pd.notna(row.iloc[7]) else '',
            description=row.iloc[8] if pd.notna(row.iloc[8]) else '',
            dating=row.iloc[9] if pd.notna(row.iloc[9]) else '',
            openatlas_class='human_remains'
            if row.iloc[6] == 'menschl. Kn.' else 'artifact')
        finds_.append(entry_obj)
    return finds_


def parse_katalog(file_path: Path) -> dict[str, FundEntity]:
    doc = Document(str(file_path))
    entities: dict[str, FundEntity] = {}
    current: Optional[FundEntity] = None

    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            first_text = row.cells[0].text.strip()
            row_id = None
            if (first_text
                    and first_text[0].isdigit() and first_text.endswith('.')):
                row_id = first_text.rstrip('.')
            is_new_entity = False
            if row_id:
                if current is None or current.id != row_id:
                    is_new_entity = True
            else:
                continue
            if is_new_entity:
                current = FundEntity(id=row_id)
                entities[row_id] = current
            if not current:
                continue
            row_seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in row_seen_tc:
                    continue
                row_seen_tc.add(tc_id)
                text = cell.text.strip()
                if text == f'{current.id}.':
                    continue
                current.add_entry(text)
    return entities


def parse_find_entry(entry: str) -> dict[str, str]:
    fndnr_match = re.search(r'Fund-Nr\.\s*([\d/]+)', entry)
    se_match = re.search(r'SE\s*(\d+)', entry)
    id_match = re.search(r'ID\s*(\d+)', entry)

    return {
        'fndnr': fndnr_match.group(1) if fndnr_match else '',
        'se': se_match.group(1) if se_match else '',
        'id': id_match.group(1) if id_match else ''}


def get_fundkatalog_entries(
        fund_entities: dict[str, FundEntity]) -> list[KatalogEntity]:
    result_ = []
    for entry in fund_entities.values():
        start_date = ''
        end_date = ''
        weight = ''
        length = ''
        height = ''
        diameter = ''
        coin_ = ''
        description_ = ''
        if entry.date:
            date = entry.date.split('\u2013')  # en dash, not hyphen!
            start_date = date[0]
            if len(date) > 1:
                end_date = date[1]
        if entry.dimensions:
            dim_tmp = entry.dimensions.split('mm')
            weight_dimensions = dim_tmp[0].split('g,')
            weight = weight_dimensions[0].strip()
            if len(weight_dimensions) == 2:
                tmp = weight_dimensions[1]
                coin_dimensions = tmp.split('h,')
                if len(coin_dimensions) == 2:
                    coin_ = coin_dimensions[0].strip()
                    tmp = coin_dimensions[1].strip()
                dimensions = tmp.split('x')
                if len(dimensions) == 2:
                    length = dimensions[0].strip()
                    height = dimensions[1].strip()
                else:
                    diameter = dimensions[0].strip()
            if len(dim_tmp) == 2:
                description_ = f'{dim_tmp[1].strip()}\n'

        # todo: check out fundnummer with subcategories e.g. 1489/4. Add
        #  for each a new artifact?

        idendification = parse_find_entry(entry.identifications)
        result_.append(KatalogEntity(
            id=entry.id,
            type=entry.type or 'Medaille',
            start_date=start_date,
            end_date=end_date,
            location=entry.location,
            description=f'{description_}{entry.description}',
            weight=weight.replace(',', '.'),
            coin=coin_.replace(',', '.'),
            length=length.replace(',', '.'),
            height=height.replace(',', '.'),
            diameter=diameter.replace(',', '.'),
            fndnr=idendification['fndnr'],
            se=idendification['se'],
            image_id=idendification['id']))
    return result_


def parse_individual_docx(file_path: Path) -> Optional[Individual]:
    doc = Document(str(file_path))
    tables = doc.tables
    if len(tables) < 4:
        return None

    # --- Table 2: contains SE id and general attributes ---
    t2 = tables[1]
    se_id = None
    for token in t2._cells:
        if token.text.startswith("SE"):
            try:
                se_id = int(
                    token.text.replace("SE", "").replace(":", "").strip())
                break
            except ValueError:
                continue
    if se_id is None:
        return None

    # --- Table 4: contains main archaeological info ---
    t4 = tables[3]
    fields = {
        "Lage": "",
        "Orientierung": "",
        "Erhaltungszustand": "",
        "Dislozierung": "",
        "Alter": "",
        "Bergung": "",
        "Art": ""}

    for row in t4.rows:
        cells = row.cells
        if len(cells) < 1:
            continue
        for cell in cells:
            text = cell.text.strip()
            splited_text = text.split(':')
            label = splited_text[0].strip()
            value = splited_text[1].strip() if len(splited_text) > 1 else ""
            if label in fields:
                fields[label] = value

    description_parts: list[str] = []
    desc_labels = {
        "Position(Objekt)+Beschreibung",
        "Grabkonstruktion",
        "Fundmaterial",
        "Grabmarkierung/ -überbau und -form",
        "Anmerkungen/ Skizze"}

    for row in t4.rows:
        cells = row.cells
        if len(cells) < 1:
            continue
        splited_text = cells[0].text.split(':')
        label = splited_text[0].strip()
        value = splited_text[1].strip() if len(splited_text) > 1 else ""
        if label in desc_labels and value and value != '-':
            description_parts.append(f"{label}: {value}")

    ind = Individual(
        se_id=se_id,
        description=description_parts,
        probe_type=fields["Art"],
        position=fields["Lage"],
        orientation=fields["Orientierung"],
        preservation=fields["Erhaltungszustand"],
        dislocation=fields["Dislozierung"],
        age=fields["Alter"],
        extraction=fields["Bergung"])
    return ind


def get_individuals() -> list[Individual]:
    output: list[Individual] = []
    for file1 in sorted(SKELETON_PATH.glob("Ind_*_SE_*.docx")):
        ind = parse_individual_docx(file1)
        if ind:
            output.append(ind)
    for file2 in sorted(SKELETON_PATH.glob("SE_*_Ind_*.docx")):
        ind = parse_individual_docx(file2)
        if ind:
            output.append(ind)
    return output


def merge_units(
        strat_units: list[ParsedStratigraphicUnit],
        individuals: list[Individual]) -> list[StratigraphicUnit]:
    individual_lookup = {ind.se_id: ind for ind in individuals}
    merged = []
    for strat_unit in strat_units:
        ind = individual_lookup.get(int(strat_unit.se_id))
        merged.append(StratigraphicUnit(
            id_=strat_unit.id_,
            se_id=strat_unit.se_id,
            name=strat_unit.name,
            layer=strat_unit.layer,
            feature=strat_unit.feature,
            type=strat_unit.type,
            individual_id=strat_unit.individual_id,
            description=ind.description if ind else [],
            probe_type=ind.probe_type if ind else None,
            # find_type=ind.find_type if ind else None,
            position=ind.position if ind else None,
            orientation=ind.orientation if ind else None,
            preservation=ind.preservation if ind else None,
            dislocation=ind.dislocation if ind else None,
            age=ind.age if ind else None,
            extraction=ind.extraction if ind else None))
    return merged


def build_types(
        entities: Sequence[Feature | StratigraphicUnit | Find],
        hierarchy: Entity,
        attribute: str) -> dict[str, Entity]:
    types: dict[str, Entity] = {}
    for entry_ in entities:
        key_ = getattr(entry_, attribute, None)
        if not key_ or key_ in types:
            continue
        # type_ = Entity.insert('type', key_)
        type_ = insert({'name': key_, 'openatlas_class_name': 'type'})
        type_.link('P127', hierarchy)
        types[key_] = type_
    return types


def build_find_types(
        entities: list[Find],
        attribute: str) -> dict[str, Entity]:
    types: dict[str, Entity] = {}
    for entry_ in entities:
        hierarchy = import_artifact_type
        if entry_.material == 'menschl. Kn.':
            hierarchy = import_hr_type
        key_ = getattr(entry_, attribute, None)
        if not key_ or key_ in types:
            continue
        # type_ = Entity.insert('type', key_)
        type_ = insert({'name': key_, 'openatlas_class_name': 'type'})
        type_.link('P127', hierarchy)
        types[key_] = type_
    return types


with (app.test_request_context()):
    app.preprocess_request()
    exact_match = get_exact_match()
    case_study = Entity.get_by_id(16305)
    folder_map = build_se_ind_map(SCHNITTE_PATH)
    cc_by_sa_type = Entity.get_by_id(50)
    file_excavation_type = Entity.get_by_id(19068)
    file_skeletonman_type = Entity.get_by_id(19069)
    file_find_type = Entity.get_by_id(19072)

    # Skelettmännchen
    skelett_maenchens = {}
    for file_ in SKELETON_IMAGE_PATH.iterdir():
        skelett_maenchens[file_.stem] = file_

    skelett_maenchens_files = {}
    for k, v in skelett_maenchens.items():
        key = f"SE_{k.split()[1].split('_')[0]}"
        if key not in skelett_maenchens_files:
            skelett_maenchens_files[key] = v

    # Fundfotos
    fundfotos = {}
    for file_ in FUNDFOTOS.iterdir():
        fundfotos[file_.stem] = file_

    # Fundfotos Medaillien
    fundfotos_medaillien = {}
    for file_ in FUNDFOTOS_MEDAILLIEN.iterdir():
        fundfotos_medaillien[file_.stem] = file_

    place = Entity.get_by_id(145)
    feature_main_type = Entity.get_by_id(72)
    import_feature_type = insert(
        {'name': 'imported', 'openatlas_class_name': 'type'})
    import_feature_type.link('P127', feature_main_type)
    artifact_main_type = Entity.get_by_id(21)
    import_artifact_type = insert(
        {'name': 'imported', 'openatlas_class_name': 'type'})
    import_artifact_type.link('P127', artifact_main_type)
    hr_main_type = Entity.get_by_id(78)
    import_hr_type = insert(
        {'name': 'imported', 'openatlas_class_name': 'type'})
    import_hr_type.link('P127', hr_main_type)
    su_main_type = Entity.get_by_id(75)
    import_su_type = insert(
        {'name': 'imported', 'openatlas_class_name': 'type'})
    import_su_type.link('P127', su_main_type)

    # Get OpenAtlas reference systems
    ref_sys_finds = Entity.get_by_id(16307)
    ref_sys_indi = Entity.get_by_id(321)
    ref_sys_feat = Entity.get_by_id(16306)
    ref_sys_su = Entity.get_by_id(16308)

    # Get OpenAtlas hierarchies
    probe_hierarchy = Entity.get_by_id(16309)
    find_hierarchy = Entity.get_by_id(16310)
    position_hierarchy = Entity.get_by_id(16311)
    orientation_hierarchy = Entity.get_by_id(16312)
    preservation_hierarchy = Entity.get_by_id(16313)
    dislocation_hierarchy = Entity.get_by_id(16314)
    age_hierarchy = Entity.get_by_id(16315)
    extraction_hierarchy = Entity.get_by_id(16316)
    layer_hierarchy = Entity.get_by_id(16317)
    material_hierarchy = Entity.get_by_id(16318)
    # designation_hierarchy = Entity.get_by_id(16319)
    dating_hierarchy = Entity.get_by_id(16320)
    cut_hierarchy = Entity.get_by_id(16323)

    # For finds
    additional_find_hierarchy = Entity.get_by_id(19207)
    coin_orientation_type = Entity.get_by_id(19206)
    diameter_type = Entity.get_by_id(19205)
    height_type = Entity.get_by_id(19202)
    weight_type = Entity.get_by_id(19204)
    length_type = Entity.get_by_id(19203)

    # Get data out of documents
    features = parse_features()
    merged_units = merge_units(
        parse_stratigraphic_units(),
        get_individuals())
    finds = parse_finds()

    # Build type dictionaries from Feature
    feature_types = build_types(features, import_feature_type, 'type')
    cut_types = build_types(features, cut_hierarchy, 'cut')

    # Build type dictionaries from StratigraphicUnit
    su_types = build_types(merged_units, import_su_type, 'type')
    probe_types = build_types(merged_units, probe_hierarchy, 'probe_type')
    find_types = build_types(merged_units, find_hierarchy, 'find_type')
    position_types = build_types(merged_units, position_hierarchy, 'position')
    age_types = build_types(merged_units, age_hierarchy, 'age')
    layer_types = build_types(merged_units, layer_hierarchy, 'layer')
    orientation_types = build_types(
        merged_units, orientation_hierarchy, 'orientation')
    preservation_types = build_types(
        merged_units, preservation_hierarchy, 'preservation')
    dislocation_types = build_types(
        merged_units, dislocation_hierarchy, 'dislocation')
    extraction_types = build_types(
        merged_units, extraction_hierarchy, 'extraction')

    # Build type dictionaries from Find
    material_types = build_types(finds, material_hierarchy, 'material')
    dating_types = build_types(finds, dating_hierarchy, 'dating')

    designation_types = build_find_types(finds, 'designation')

    added_features: dict[str, Entity] = {}
    for fe_entry in features:
        feature = insert({
            'name': fe_entry.name,
            'description': fe_entry.description,
            'openatlas_class_name': 'feature'})
        feature.link('P2', case_study)
        feature.link('P2', feature_types[fe_entry.type])
        if fe_entry.cut:
            feature.link('P2', cut_types[fe_entry.cut])
        feature.link('P46', place, inverse=True)
        ref_sys_feat.link(
            'P67',
            feature,
            str(fe_entry.obj_id),
            type_id=exact_match.id)
        location = insert({
            'name': f'Location of {fe_entry.name}',
            'openatlas_class_name': 'object_location'})
        feature.link('P53', location)
        added_features[fe_entry.id_] = feature

    added_stratigraphic: dict[str, Entity] = {}
    for entry in merged_units:
        desc = '\n'.join(entry.description) if entry.description else ''
        su = insert({
            'name': entry.name,
            'description': desc,
            'openatlas_class_name': 'stratigraphic_unit'})
        su.link('P2', case_study)
        su.link('P46', added_features[entry.feature], inverse=True)
        ref_sys_su.link(
            'P67',
            su,
            str(entry.se_id),
            type_id=exact_match.id)
        if entry.individual_id:
            ref_sys_indi.link(
                'P67',
                su,
                str(entry.individual_id),
                type_id=exact_match.id)
        location = insert({
            'name': f'Location of {entry.name}',
            'openatlas_class_name': 'object_location'})
        su.link('P53', location)

        if entry.type:
            su.link('P2', su_types[entry.type])
        if entry.probe_type:
            su.link('P2', probe_types[entry.probe_type])
        if entry.find_type:
            su.link('P2', find_types[entry.find_type])
        if entry.position:
            su.link('P2', position_types[entry.position])
        if entry.orientation:
            su.link('P2', orientation_types[entry.orientation])
        if entry.preservation:
            su.link('P2', preservation_types[entry.preservation])
        if entry.dislocation:
            su.link('P2', dislocation_types[entry.dislocation])
        if entry.age:
            su.link('P2', age_types[entry.age])
        if entry.extraction:
            su.link('P2', extraction_types[entry.extraction])
        if entry.layer:
            su.link('P2', layer_types[entry.layer])

        skelett_file = skelett_maenchens_files.get(f'SE_{entry.se_id}')
        if skelett_file:
            file = insert({
                'name': f'SE {entry.se_id} Ind '
                        f'{entry.individual_id} Skelettmännchen',
                'openatlas_class_name': 'file'})
            file.save_file_info(FILE_INFO)
            file.link('P2', cc_by_sa_type)
            file.link('P2', file_skeletonman_type)
            su.link('P67', file, inverse=True)

            ext = skelett_file.suffix
            dest = UPLOAD / f"{file.id}{ext}"
            UPLOAD.mkdir(parents=True, exist_ok=True)
            shutil.copy(skelett_file, dest)

        key = f"SE_{entry.se_id}_Ind_{entry.individual_id}"

        if folder := folder_map.get(key):
            images = [
                p for p in folder.iterdir()
                if p.is_file() and p.suffix.lower() in {".jpg", ".jpeg"}]
            for idx, image in enumerate(images):
                file = insert({
                    'name': f'SE {entry.se_id} Ind '
                            f'{entry.individual_id} {idx}',
                    'openatlas_class_name': 'file'})
                file.save_file_info(FILE_INFO)
                file.link('P2', cc_by_sa_type)
                file.link('P2', file_excavation_type)
                su.link('P67', file, inverse=True)
                ext = image.suffix
                dest = UPLOAD / f"{file.id}{ext.lower()}"
                UPLOAD.mkdir(parents=True, exist_ok=True)
                shutil.copy(image, dest)

        added_stratigraphic[entry.id_] = su

    added_finds: dict[str, Entity] = {}
    for fi_entry in finds:
        if fi_entry.stratigraphic_unit:
            link_to_entity = added_stratigraphic.get(
                fi_entry.stratigraphic_unit)
        elif fi_entry.feature_id:
            link_to_entity = added_features.get(fi_entry.feature_id)
        else:
            link_to_entity = place
        if not link_to_entity:
            DEBUG_MSG['no_super_for_finds'].append(fi_entry.f_id)
            continue

        system_class = 'artifact'
        if fi_entry.material == 'menschl. Kn.':
            system_class = 'human_remains'
        find = insert({
            'name': fi_entry.name,
            'description': fi_entry.description,
            'openatlas_class_name': system_class})
        find.link('P2', case_study)
        find.link('P46', link_to_entity, inverse=True)

        ref_sys_finds.link(
            'P67',
            find,
            str(fi_entry.f_id),
            type_id=exact_match.id)
        location = insert({
            'name': f'Location of {fi_entry.name}',
            'openatlas_class_name': 'object_location'})
        find.link('P53', location)

        for fund_foto, path_ in fundfotos.items():
            fund_foto_name = fund_foto.split('_')[0]

            if fund_foto_name == fi_entry.name:
                fundfoto_file = path_
                file = insert({
                    'name': f'{fi_entry.name}',
                    'openatlas_class_name': 'file'})
                file.save_file_info(FILE_INFO)
                file.link('P2', cc_by_sa_type)
                file.link('P2', file_find_type)
                find.link('P67', file, inverse=True)

                ext = fundfoto_file.suffix
                dest = UPLOAD / f"{file.id}{ext}"
                UPLOAD.mkdir(parents=True, exist_ok=True)
                shutil.copy(fundfoto_file, dest)

        if fi_entry.material:
            find.link('P2', material_types[fi_entry.material])
        if fi_entry.designation:
            find.link('P2', designation_types[fi_entry.designation])
        if fi_entry.dating:
            find.link('P2', dating_types[fi_entry.dating])
        added_finds[str(fi_entry.f_id)] = find

    if FUNDFOTOS_KATALOG.exists():
        result: dict[str, FundEntity] = parse_katalog(FUNDFOTOS_KATALOG)
        fundkatalog_entries = get_fundkatalog_entries(result)
        for fundkatalog_entry in fundkatalog_entries:
            find_entity = added_finds[fundkatalog_entry.fndnr.split('/')[0]]
            # todo: add count for name, so it should be name_1, name_2 etc.
            new_name = find_entity.name
            inser_data = {
                'name': new_name,
                'description': fundkatalog_entry.description.strip(),
                'openatlas_class_name': 'artifact'}
            if fundkatalog_entry.start_date:
                inser_data.update({
                    'begin_from': f'{fundkatalog_entry.start_date}-01-01',
                    'begin_to': f'{fundkatalog_entry.start_date}-12-01'})
            if fundkatalog_entry.end_date:
                inser_data.update({
                    'end_from': f'{fundkatalog_entry.end_date}-01-01',
                    'end_to': f'{fundkatalog_entry.end_date}-12-01'})
            additional_find = insert(inser_data)

            additional_find.link('P2', case_study)
            additional_find.link('P46', find_entity, inverse=True)
            location = insert({
                'name': f'Location of {new_name}',
                'openatlas_class_name': 'object_location'})

            if fundkatalog_entry.weight:
                additional_find.link(
                    'P2',
                    weight_type,
                    description=fundkatalog_entry.weight)
            if fundkatalog_entry.diameter:
                additional_find.link(
                    'P2',
                    diameter_type,
                    description=fundkatalog_entry.diameter)
            if fundkatalog_entry.coin:
                additional_find.link(
                    'P2',
                    coin_orientation_type,
                    description=fundkatalog_entry.coin)
            if fundkatalog_entry.length:
                additional_find.link(
                    'P2',
                    length_type,
                    description=fundkatalog_entry.length)
            if fundkatalog_entry.height:
                additional_find.link(
                    'P2',
                    height_type,
                    description=fundkatalog_entry.height)
            for fund_foto, path_ in fundfotos_medaillien.items():
                fund_foto_name = fund_foto.split('_')[0]

                if fund_foto_name == fundkatalog_entry.image_id:
                    fundfoto_file = path_
                    file = insert({
                        'name': f'{fi_entry.name} {fundkatalog_entry.image_id}',
                        'openatlas_class_name': 'file'})
                    file.save_file_info(FILE_INFO)
                    file.link('P2', cc_by_sa_type)
                    file.link('P2', file_find_type)
                    additional_find.link('P67', file, inverse=True)

                    ext = fundfoto_file.suffix
                    dest = UPLOAD / f"{file.id}{ext}"
                    UPLOAD.mkdir(parents=True, exist_ok=True)
                    shutil.copy(fundfoto_file, dest)
    else:
        print(f'Datei nicht gefunden: {FUNDFOTOS_KATALOG}')
    print(DEBUG_MSG)
